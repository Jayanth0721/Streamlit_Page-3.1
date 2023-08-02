import streamlit as st
from PIL import Image
from pydub import AudioSegment
import io
import subprocess
import os
import comtypes.client
import tempfile
import pythoncom
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import pythoncom
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Adjust the layout width of Streamlit
st.set_page_config(layout="wide")

st.toast('Welcome to 3.1', icon='😍')

# Print the text "Welcome to" in black color and "3.1" in dark violet color
st.markdown('<p style="font-size: 38px; color: black; display: inline;">&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp;&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp;&nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; &nbsp; Welcome to </p>'
            '<p style="font-size: 52px; color: #9400D3; display: inline;">"3.1"</p>',
            unsafe_allow_html=True)
# Path to the image file
image_path = "C:\\de\\3.png"

# Display the image to fit the screen width
try:
    img = Image.open(image_path)
    st.image(img, caption='3.1', use_column_width=True)
except FileNotFoundError:
    st.error("Image file not found. Please check the file path.")
except Exception as e:
    st.error(f"Error occurred while displaying the image: {e}")

st.markdown("**Features of 3.1 (3 in 1)**")
st.markdown("1.Media Player")
st.write('--You Play Audio, Video & You Can View Images.')
st.markdown(" ")
st.markdown("2.Programming Assistant.")
st.write("--You Can Get Some Basic Program Code Using Python howdoi Library.")
st.markdown(" ")
st.markdown("3.File Convertor.")
st.write("--You Can Convert PDF To Document, Document To PDF. ")

st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")

st.title('Media Player')

def play_audio(audio_file):
    # Play the audio file using the st.audio function
    st.audio(audio_file, format="audio/" + audio_file.type.split("/")[-1])

def play_video(video_file):
    # Play the video file using the embedded HTML5 video player
    st.video(video_file, format="video/" + video_file.type.split("/")[-1])

def display_image(image_file):
    # Display the image using st.image
    st.image(image_file, caption=image_file.name, use_column_width=True)

# Prompt the user to choose the file format
file_format = st.selectbox("Select a file format", ["Audio", "Video", "Image"])

# Depending on the selected file format, display the corresponding file uploader and viewer
if file_format == "Audio":
    # Get the audio file path
    audio_file = st.file_uploader("Upload an audio file", type=["mp3", "wav"])
    if audio_file is not None:
        play_audio(audio_file)
elif file_format == "Video":
    # Get the video file path
    video_file = st.file_uploader("Upload a video file", type=["mp4", "avi", "MKV"])
    if video_file is not None:
        # Custom text disclaimer
        st.markdown("<p style='font-size: 14px; color: red; text-align: center;'>Video is paused. Use the right arrow key to forward 15 seconds and the left arrow key to go back 15 seconds.</p>", unsafe_allow_html=True)
        play_video(video_file)
elif file_format == "Image":
    # Get the image file path
    image_file = st.file_uploader("Upload an image file", type=["jpg", "png", "jpeg"])
    if image_file is not None:
        display_image(image_file)

 
 
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")

#howdoi 
def get_answer(question):
    """Gets an answer to the given question using howdoi."""
    try:
        command = f'howdoi "{question}"'
        result = subprocess.run(command, capture_output=True, text=True, shell=True)
        answer = result.stdout.strip()
        return answer
    except Exception as e:
        return f"An error occurred: {str(e)}"

def main():
    """The main function of the app."""
    st.title("Ask a Programming Question using howdoi")

    # Input: User's question
    question = st.text_input("Enter a Programming Question:")
    if question:
        st.write("You asked: ", question)

        # Get the answer using howdoi
        try:
            answer = get_answer(question)
            st.subheader("Answer:")
            st.code(answer, language='python')
        except Exception as e:
            st.error("An error occurred while fetching the answer.")
            st.error(str(e))

if __name__ == "__main__":
    main()     


st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")
st.markdown(" ")


#File Convertor


def set_heading_style(paragraph):
    run = paragraph.runs[0]
    run.font.size = Pt(20)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.space_after = Pt(12)

def docx_to_pdf(docx_path):
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    # Initialize COM library in the main thread
    pythoncom.CoInitialize()

    # Initialize COM object for Word application
    word = comtypes.client.CreateObject("Word.Application")

    try:
        # Open the DOCX file
        doc = word.Documents.Open(docx_path)

        # Save as PDF
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
    except Exception as e:
        st.error(f"Error converting file: {e}")
    finally:
        # Close the Word application
        word.Quit()

    # Uninitialize the COM library
    pythoncom.CoUninitialize()

    return pdf_path

def pdf_to_docx(pdf_path):
    docx_path = os.path.splitext(pdf_path)[0] + ".docx"

    # Extract text from PDF
    pdf_text = ""
    with open(pdf_path, "rb") as f:
        pdf_reader = PdfReader(f)
        for page in pdf_reader.pages:
            pdf_text += page.extract_text()

    # Create DOCX file and write the extracted text
    doc = Document()
    title = doc.add_paragraph("Converted Document", style="Heading 1")
    set_heading_style(title)

    p = doc.add_paragraph(pdf_text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    try:
        # Save as DOCX
        doc.save(docx_path)
    except Exception as e:
        st.error(f"Error converting file: {e}")

    return docx_path

def main():
    st.title("File Format Converter")

    # File upload
    file = st.file_uploader("Upload a file", type=["pdf", "docx"])

    # Convert and display download link
    if file is not None:
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(file.read())

        if file.type == "application/pdf":
            output_file = pdf_to_docx(tmp_file.name)
        else:
            output_file = docx_to_pdf(tmp_file.name)

        st.success("File converted successfully!")

        with open(output_file, "rb") as f:
            st.download_button(label="Download Converted File", data=f, file_name=os.path.basename(output_file))

        os.remove(tmp_file.name)
        os.remove(output_file)

if __name__ == "__main__":
    main()

st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("""
        <div style="text-align: center;">
            <h3>Thanks For Using 3.1(3 in 1)!</h3>
        </div>
    """, unsafe_allow_html=True)